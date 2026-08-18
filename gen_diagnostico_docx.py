from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

NAVY  = RGBColor(0x0C, 0x23, 0x40)
NAVY2 = RGBColor(0x1A, 0x3A, 0x5C)
GOLD  = RGBColor(0xB8, 0x94, 0x3A)
GREEN = RGBColor(0x1D, 0x6B, 0x4F)
RED   = RGBColor(0xB8, 0x32, 0x32)
SLATE = RGBColor(0x7A, 0x92, 0xA8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_text(row, values, color=None, bold=False, size=10, bg=None):
    if bg:
        shade_row(row, bg)
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(size)
                run.bold = bold
                if color:
                    run.font.color.rgb = color

def add_heading(doc, text, size=13, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

# ── HEADER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FIBRA HOTEL CORPORATIVO  ·  ÁREA DE ARRENDAMIENTO  ·  OBJETIVO 2")
r.font.size = Pt(9); r.font.color.rgb = GOLD; r.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("DIAGNÓSTICO GLOBAL DE ENERGÉTICOS ARRENDATARIOS 2026")
r2.font.size = Pt(18); r2.bold = True; r2.font.color.rgb = NAVY

doc.add_paragraph()

meta_t = doc.add_table(rows=1, cols=3)
meta_t.style = 'Table Grid'
set_row_text(meta_t.rows[0],
    ["Elaborado por: Gerardo Vega Quinto", "Período: Enero – Julio 2026", "Fecha: 25 de mayo de 2026"],
    color=WHITE, bold=True, size=10, bg="0C2340")
for cell in meta_t.rows[0].cells:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── RESUMEN EJECUTIVO ──
add_heading(doc, "RESUMEN EJECUTIVO", size=12)

res_t = doc.add_table(rows=2, cols=4)
res_t.style = 'Table Grid'
set_row_text(res_t.rows[0],
    ["Total arrendatarios", "Independizados", "Dependientes del hotel", "Propiedades analizadas"],
    color=WHITE, bold=True, size=10, bg="0C2340")
set_row_text(res_t.rows[1],
    ["53  (31 locales · 22 antenas)", "35  (66.0%)", "18  (34.0%)", "29 hoteles"],
    color=WHITE, bold=True, size=11, bg="1A3A5C")

doc.add_paragraph()

# ── LOCALES ──
add_heading(doc, "LOCALES COMERCIALES — Requieren independencia de electricidad y agua")

doc.add_paragraph("Los locales comerciales deben contar con medidor propio de electricidad y toma de agua independiente del sistema hotelero. A continuación se detalla el estatus por hotel.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

local_data = [
    ("1ACCM", "One Acapulco Costera", [
        ("TV Azteca SAB de CV", "Local TV Azteca", "Indep.", "Indep.", "Independizado"),
        ("Circle K (★ Nuevo Q1-26)", "Local #2", "Indep.", "Indep.", "Sin reembolso registrado"),
    ]),
    ("1AGCI", "One Aguascalientes Sur", [
        ("Comercial Automotriz de los Altos", "Local Automotriz", "Indep.", "Indep.", "Independizado"),
    ]),
    ("1MTAP", "One Monterrey Aeropuerto", [
        ("7-Eleven México SA de CV", "Local 7-Eleven", "Indep.", "Indep.", "Independizado"),
        ("Servicios Inmobiliarios Alsea SA de CV", "Local Starbucks", "Indep.", "Indep.", "Independizado"),
        ("Consolidadora de Servicios Aduanales SC", "Local 4B", "Indep.", "Indep.", "Independizado"),
        ("Primeflight Aviation Services México", "Bodega Aerocharter", "Indep.", "Indep.", "Independizado"),
        ("Gloria JH S de RL de CV", "Local Gloria", "Indep.", "Reembolso hotel", "DEPENDIENTE — agua"),
        ("Proyectos Electronicos de Entretenimiento", "Restaurante Hole in One", "Indep.", "Reembolso hotel", "DEPENDIENTE — agua"),
    ]),
    ("1PUFI", "One Puebla Finsa", [
        ("Agro-Partes Puebla", "Local Agro-Partes", "Indep.", "Indep.", "Independizado"),
    ]),
    ("1XCUA", "One Cuautitlan", [
        ("Banco Santander México SA", "Local Santander", "Indep.", "Indep.", "Independizado"),
        ("Waldo's Dolar Mart de México", "Local Waldo's", "Indep.", "Indep.", "Independizado"),
        ("Zapata", "Local Ford", "Indep.", "Indep.", "Independizado"),
    ]),
    ("FAAG", "Fiesta Americana Aguascalientes", [
        ("David Valdez Romo", "Local Fishers", "Indep.", "Indep.", "Independizado"),
    ]),
    ("FACC", "Fiesta Americana Condesa Cancun", [
        ("Amazon Tours", "Local Aqua World", "Indep.", "Indep.", "Independizado"),
        ("T.I. Experiencias en Grupos y Convenciones", "Local Tropical Elite Travel", "Indep.", "Indep.", "Independizado"),
    ]),
    ("FAHE", "Fiesta Americana Hermosillo", [
        ("Operadora Valiente", "Sonora Grill", "Reembolso hotel", "Reembolso hotel", "DEPENDIENTE — luz y agua"),
    ]),
    ("FICDJ", "Fiesta Inn Ciudad Juarez", [
        ("Maria Irene Vota Echavarri", "Tabaqueria", "Reembolso hotel", "N/A", "DEPENDIENTE — electricidad"),
    ]),
    ("FICDO", "Fiesta Inn Ciudad Obregon", [
        ("Super Carnes Clasificadas Chihuahua", "Local Carniceria", "Indep.", "Indep.", "Independizado"),
        ("Coyotas Lulu SA de CV", "Local Reposteria", "Reembolso hotel", "N/A", "DEPENDIENTE — electricidad"),
    ]),
    ("FIPER", "Fiesta Inn Perinorte", [
        ("DIVOL SA", "Local DIVOL", "Indep.", "Indep.", "Independizado"),
    ]),
    ("FIQRO", "Fiesta Inn Queretaro", [
        ("Click Mobility", "Local Alamo", "Indep.", "Indep.", "Independizado"),
    ]),
    ("GTIOT", "Gamma Tijuana Otay", [
        ("DAQU de Sonora SA de CV", "Dairy Queen", "Reembolso hotel", "Reembolso hotel", "DEPENDIENTE — luz y agua"),
        ("Yokomy Sushi S de RL de CV", "Local Yokomi Sushi", "Reembolso hotel", "Reembolso hotel", "DEPENDIENTE — luz y agua"),
        ("Julio Cesar Torres Rojo", "Renta Local", "Reembolso hotel", "Reembolso hotel", "DEPENDIENTE — luz y agua"),
    ]),
    ("MTYCW", "Wyndham Ambassador Monterrey Centro", [
        ("Alquiladora de Vehículos Automotores", "Local Hertz", "Indep.", "Indep.", "Independizado"),
        ("Click Mobility", "Local Alamo", "Indep.", "Indep.", "Independizado"),
        ("Auto Arrendadora y Promotora SA de CV", "Local Sixt", "Indep.", "Indep.", "Independizado"),
        ("Juan Carlos Alberto Ayala Olvera", "Local Oficina", "Indep.", "Indep.", "Independizado"),
        ("Corporativo Empresarial Solar", "Local Corporativo", "Indep.", "Indep.", "Independizado"),
        ("Gestoria Administrativa Profesional", "Local Gestoria", "Indep.", "Indep.", "Independizado"),
    ]),
]

for code, hotel_name, tenants in local_data:
    has_dep = any("DEPENDIENTE" in t[4] for t in tenants)
    p = doc.add_paragraph()
    r = p.add_run(f"{code}  ·  {hotel_name}")
    r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = RED if has_dep else GREEN

    n_rows = 1 + len(tenants)
    tbl = doc.add_table(rows=n_rows, cols=5)
    tbl.style = 'Table Grid'
    set_row_text(tbl.rows[0], ["Arrendatario", "Local", "Electricidad", "Agua", "Estatus"],
                 color=WHITE, bold=True, size=9, bg="0C2340")
    for i, (name, loc, luz, agua, est) in enumerate(tenants):
        is_dep = "DEPENDIENTE" in est
        bg = "FDEAEA" if is_dep else "FFFFFF"
        set_row_text(tbl.rows[i+1], [name, loc, luz, agua, est], size=9, bg=bg)
        if is_dep:
            for cell in tbl.rows[i+1].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RED

    doc.add_paragraph()

# ── ANTENAS ──
add_heading(doc, "ANTENAS — Solo requieren independencia de electricidad")

dep_antenas = [
    ("Fiesta Americana Condesa Cancun", "FACC", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Fiesta Americana Hermosillo", "FAHE", "Intelli Site Solutions SAPI de CV"),
    ("Fiesta Americana Hacienda Galindo", "FAHG", "Grupo AT&T Celullar"),
    ("Fiesta Inn Cuautitlan", "FICUA", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Fiesta Inn Durango", "FIDUR", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Fiesta Inn Morelia Altozano", "FIMOA", "MXT Capital Partners SAPI de CV"),
    ("Fiesta Inn Perinorte", "FIPER", "TXM Terrenos S de RL de CV"),
    ("Fiesta Inn Perisur", "FISUR", "Megacable Comunicaciones de Mexico SA de CV"),
    ("One Monterrey Aeropuerto", "1MTAP", "Megacable Comunicaciones de Mexico SA de CV"),
    ("One Toluca Aeropuerto", "1TOAP", "TXM Terrenos S de RL de CV"),
]

ind_antenas = [
    ("One Guadalajara Tapatio", "1GDTA", "Nuviax SA de CV"),
    ("AC Antea Queretaro", "QROQA", "MXT Capital Partners SAPI de CV"),
    ("Camino Real Puebla", "CRPUE", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Camino Real Puebla", "CRPUE", "ATyT Comunicaciones Digitales S de RL de CV"),
    ("Fiesta Americana Hacienda Galindo", "FAHG", "Operadora de Sites Mexicanos SA de CV"),
    ("Fiesta Inn Chihuahua", "FICHI", "ATyT Comunicaciones Digitales S de RL de CV"),
    ("Fiesta Inn Chihuahua", "FICHI", "Intelli Site Solutions SAPI de CV"),
    ("Fiesta Inn Guadalajara Expo", "FIGDL", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Fiesta Inn Nuevo Laredo", "FINVL", "ATyT Comunicaciones Digitales S de RL de CV"),
    ("Fiesta Inn Oaxaca", "FIOAX", "Canalizaciones y Accesos Profesionales S de RL"),
    ("Fiesta Inn Perinorte", "FIPER", "Mexico Tower Partners SAPI de CV"),
    ("Live Aqua San Miguel de Allende", "AQSM", "Mexico Tower Partners"),
]

p = doc.add_paragraph()
r = p.add_run("Antenas con electricidad dependiente del hotel (10)")
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RED

dep_tbl = doc.add_table(rows=1+len(dep_antenas), cols=3)
dep_tbl.style = 'Table Grid'
set_row_text(dep_tbl.rows[0], ["Hotel", "Clave", "Arrendatario (Operador)"],
             color=WHITE, bold=True, size=9, bg="B83232")
for i, (hotel, code, name) in enumerate(dep_antenas):
    set_row_text(dep_tbl.rows[i+1], [hotel, code, name], size=9, bg="FDEAEA")
    for cell in dep_tbl.rows[i+1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RED

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Antenas con electricidad independiente (12)")
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GREEN

ind_tbl = doc.add_table(rows=1+len(ind_antenas), cols=3)
ind_tbl.style = 'Table Grid'
set_row_text(ind_tbl.rows[0], ["Hotel", "Clave", "Arrendatario (Operador)"],
             color=WHITE, bold=True, size=9, bg="1D6B4F")
for i, (hotel, code, name) in enumerate(ind_antenas):
    set_row_text(ind_tbl.rows[i+1], [hotel, code, name], size=9, bg="E6F4EE")
    for cell in ind_tbl.rows[i+1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = GREEN

doc.add_paragraph()

# ── RESUMEN CONSOLIDADO ──
add_heading(doc, "RESUMEN CONSOLIDADO — PRIORIDADES DE INDEPENDIZACIÓN")

sum_t = doc.add_table(rows=7, cols=7)
sum_t.style = 'Table Grid'
set_row_text(sum_t.rows[0],
    ["Tipo", "Total", "Independizados", "% Indep.", "Dependientes", "% Depend.", "Prioridad"],
    color=WHITE, bold=True, size=9, bg="0C2340")

sum_data = [
    ("Locales — Luz y agua", "5", "1", "20%", "4 (Tijuana x3, Hermosillo)", "80%", "Alta"),
    ("Locales — Solo agua", "2", "0", "0%", "2 (One Monterrey)", "100%", "Alta"),
    ("Locales — Solo luz", "2", "0", "0%", "2 (Cdad. Juarez, Cdad. Obregon)", "100%", "Alta"),
    ("Locales — Totalmente independientes", "22", "22", "100%", "—", "0%", "Cumplido"),
    ("Antenas — Electricidad del hotel", "10", "0", "0%", "10 sitios en 10 hoteles", "100%", "Media"),
    ("Antenas — Electricidad independiente", "12", "12", "100%", "—", "0%", "Cumplido"),
]
bgs = ["FFFFFF", "EBF0F6", "FFFFFF", "EBF0F6", "FFFFFF", "EBF0F6"]
for i, (row_data, bg) in enumerate(zip(sum_data, bgs)):
    set_row_text(sum_t.rows[i+1], list(row_data), size=9.5, bg=bg)
    # Color priority cell
    prio_cell = sum_t.rows[i+1].cells[6]
    prio = row_data[6]
    for para in prio_cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RED if prio == "Alta" else (GREEN if prio == "Cumplido" else RGBColor(0xB8, 0x6A, 0x00))
            run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Elaborado por: Gerardo Vega Quinto — Área de Arrendamiento, Fibra Hotel Corporativo")
r.font.size = Pt(9.5); r.font.color.rgb = SLATE; r.italic = True

doc.save('/home/user/Contratos-Arrendamiento/Diagnostico_Energeticos_2026_GVega.docx')
print("DOCX OK")
